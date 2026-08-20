# -*- coding: utf-8 -*-
"""论文风 Word 生成公共工具（供 gen_product_doc.py / gen_design_doc.py 使用）

- 样式：黑体标题 / 宋体正文 / Times New Roman 西文 / 黑色 / 1.5 倍行距 / 三线表
- 公式：LaTeX 字符串 → MathML（latex2mathml）→ OMML（微软 Mml2OMML.XSL）→ 插入 Word，
  在 Word 中即显示为可读的数学公式（分数/求和/上下标，无需肉眼编译）。
"""
import latex2mathml.converter
import lxml.etree as etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 微软官方 XSL：LaTeX→MathML 后经此转 OMML（Word 数学公式 XML）
MML2OMML_XSL = r'C:\Program Files\Microsoft Office\root\Office16\Mml2OMML.XSL'
_XSLT = etree.XSLT(etree.parse(MML2OMML_XSL))

BLACK = RGBColor(0x00, 0x00, 0x00)


def _set_run(run, cn='宋体', en='Times New Roman', size=10.5, bold=False, italic=False):
    run.font.name = en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
    # 默认正文样式：宋体五号 + Times 西文
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, cn='黑体', size=16, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, cn='黑体', size=13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, cn='黑体', size=11, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(doc, text, indent=2, align=None, size=10.5, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if align:
        p.alignment = align
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(2)
    r = p.add_run(text)
    _set_run(r, size=10.5)
    return p


def formula(doc, latex):
    """LaTeX 公式 → Word 数学公式（OMML），居中段落。"""
    mathml = latex2mathml.converter.convert(latex)
    omml = str(_XSLT(etree.fromstring(mathml)))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 去掉 XML 声明
    omml = omml.split('?>', 1)[-1].strip()
    p._p.append(etree.fromstring(omml))
    return p


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """为单元格设置边框线（sz 单位 1/8pt，val=single）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge, sz in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if sz is None:
            continue
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), '000000')


def three_line_table(doc, headers, rows, widths=None):
    """三线表（booktabs 风格）：顶线 1.5pt、表头下线 0.75pt、底线 1.5pt，无竖线。"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = 1  # center
    t.autofit = False
    # 表头
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext)
        _set_run(r, cn='黑体', size=10, bold=True)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) < 14 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            _set_run(r, size=10)
    # 边框：表头行上 1.5 下 0.75；数据行只底 1.5
    for ci in range(len(headers)):
        _set_cell_borders(t.rows[0].cells[ci], top=12, bottom=6)
        for ri in range(1, len(t.rows)):
            _set_cell_borders(t.rows[ri].cells[ci], bottom=12 if ri == len(t.rows) - 1 else None)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t
