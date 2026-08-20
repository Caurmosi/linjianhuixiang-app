# -*- coding: utf-8 -*-
"""生成《林间回响》技术设计说明书.docx（论文风 + Word 数学公式）

公式为 LaTeX 字符串 → MathML → OMML，在 Word 中显示为可读数学公式
（分数 / 求和 / 上下标等真实排版，无需肉眼编译）。

以后算法有改动：编辑本脚本对应章节 → 重新运行
  python docs/gen_design_doc.py
"""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx_utils import (new_doc, h1, h2, h3, para, bullet, formula,
                        three_line_table, _set_run)


def build():
    doc = new_doc()

    # ============ 封面 ============
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run('《林间回响》'), cn='黑体', size=26, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run('城市鸟类宜居度声学诊断平台'), cn='黑体', size=22, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run('技 术 设 计 说 明 书'), cn='黑体', size=18)
    for _ in range(4):
        doc.add_paragraph()
    for line in ['版本：V1.0（2026-08）', '面向：技术评审 / 竞赛论文技术章节 / 工程维护']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p.add_run(line), size=12)
    doc.add_page_break()

    # ============ 目录 ============
    h1(doc, '目录')
    for t in ['一、系统总体架构', '二、技术栈', '三、音频采集与预处理', '四、鸟声识别算法',
              '五、生态声学指数', '六、人为噪声占比估计', '七、宜居度耦合模型', '八、时间热力图',
              '九、公共地图系统（聚合 / 隐私 / 模糊）', '十、数据看板与统计', '十一、生态简报',
              '十二、账号与云同步', '十三、前端 App 架构', '十四、公共地图网页', '十五、部署与运维',
              '十六、测试与质量保障', '十七、参考文献']:
        para(doc, t, indent=0)
    doc.add_page_break()

    # ============ 一、系统总体架构 ============
    h1(doc, '一、系统总体架构')
    h2(doc, '1.1 架构总览')
    para(doc, '系统为「端—云」双层架构：Android App（React 18 H5 壳）→ 云端后端（FastAPI，Docker 容器）→ '
              'SQLite 数据库；公共地图网页（React + MapLibre GL）独立托管，只读访问云端公共接口。')
    three_line_table(doc,
                     ['层', '技术', '职责'],
                     [['表现层 App', 'React 18 + Vite 5 + Tailwind CSS，Android WebView 套壳', '录音、导入、结果展示、地区管理、图鉴、云同步'],
                      ['表现层 网页', 'React 18 + MapLibre GL 3.6 + 高德瓦片', '公共地图浏览 / 检索 / 对比 / 看板 / 导出'],
                      ['服务层', 'FastAPI + uvicorn', 'REST API：分析、识别、公共数据、账号、备份、瓦片代理'],
                      ['算法层', 'BirdNET TFLite + numpy/scipy 自研 DSP', '鸟声识别、声学指数、噪声估计、宜居度合成'],
                      ['数据层', 'SQLite（可换 PostgreSQL）', '用户、Token、公共记录、云备份'],
                      ['部署层', 'Docker / podman（阿里云轻量服务器），Sealos 备', '容器化部署、数据卷持久化、开机自启']],
                     widths=[2.6, 4.4, 5.2])
    h2(doc, '1.2 核心数据流')
    bullet(doc, '录音分析：音频字节 → decode（ffmpeg/pyav）→ 重采样 48kHz/mono → 高通滤波 → BirdNET 推理 + STFT 指数 → 宜居度合成 → 返回 App；')
    bullet(doc, '公共上传：App 保存的地区快照 → POST /api/public/records（Bearer Token）→ 落库（cluster_key 聚合键）→ 网页实时聚合查询；')
    bullet(doc, '云同步：本地 JSON 快照 → POST /api/sync/backup（单行 upsert，≤2MB）→ 换机登录恢复。')

    # ============ 二、技术栈 ============
    h1(doc, '二、技术栈')
    three_line_table(doc,
                     ['模块', '选型', '说明'],
                     [['后端框架', 'FastAPI 0.11x + uvicorn', '异步 ASGI，自动 OpenAPI 文档，Pydantic 契约校验'],
                      ['音频解码', 'ffmpeg（外部进程）+ PyAV 双通道', 'mp3/m4a/webm 等容器解码；WAV 直读'],
                      ['数值计算', 'numpy + scipy + soundfile', 'STFT、滤波器组、指数计算；自研 DSP 不依赖 librosa'],
                      ['推理引擎', 'BirdNET GLOBAL 6K V2.4（TFLite）', '本地推理，离线可用；模型缺失自动降级启发式引擎'],
                      ['前端 App', 'React 18 + Vite 5 + Tailwind 3', 'H5 单页应用，Android WebView 套壳（Kotlin）'],
                      ['公共网页', 'React 18 + MapLibre GL 3.6 + qrcode', '高德栅格瓦片代理；二维码分享'],
                      ['数据库', 'SQLite（WAL）', '单文件、零运维；预留 PostgreSQL 迁移'],
                      ['部署', 'Docker/podman + nginx（阿里云）', '容器化、数据卷 /app/data、开机自启、HTTPS 反代'],
                      ['CI', 'GitHub Actions', '后端镜像构建推送 ghcr；Android APK 自动打包']],
                     widths=[2.8, 5.2, 4.2])

    # ============ 三、音频预处理 ============
    doc.add_page_break()
    h1(doc, '三、音频采集与预处理')
    h2(doc, '3.1 采集规格')
    para(doc, '手机麦克风录音（App 内录音最高 48kHz/16bit）；导入文件支持 wav/mp3/m4a/webm 等。'
              '后端统一重采样至 48kHz、单声道、float32。')
    h2(doc, '3.2 解码与预处理链路')
    bullet(doc, '解码：优先 PyAV（内存解码），失败回退 ffmpeg 子进程；')
    bullet(doc, '重采样：scipy.signal.resample_poly 至 48kHz；')
    bullet(doc, '单声道：多声道按均值混缩（_to_mono）；')
    bullet(doc, '高通滤波：去除 200Hz 以下交通/机械低频（削弱风噪与电机声）。')
    h2(doc, '3.3 短时傅里叶变换（STFT）')
    para(doc, '声学指数与热力图的基础为 STFT 幅度谱：')
    formula(doc, r'X(t,f)=\left|\sum_{n} x(n)\,w(n-tH)\,e^{-j2\pi fn/N}\right|')
    para(doc, '其中 x(n) 为时域采样序列，w(·) 为窗函数（Hann），N 为 FFT 长度，H 为帧移。'
              '参数：N=1024，H=512；48kHz 下频率分辨率约 46.9Hz，帧长约 21.3ms，时间分辨率约 10.7ms。')
    h2(doc, '3.4 Mel 滤波器组')
    para(doc, '识别与可视化使用 Mel 刻度，频率到 Mel 的映射为：')
    formula(doc, r'\mathrm{mel}(f)=2595\,\log_{10}\left(1+\frac{f}{700}\right)')
    para(doc, '三角滤波器组（Slaney 归一化），默认 128 个 Mel 频带，fmin=0，fmax=采样率/2。')

    # ============ 四、鸟声识别 ============
    doc.add_page_break()
    h1(doc, '四、鸟声识别算法')
    h2(doc, '4.1 主引擎：BirdNET GLOBAL 6K V2.4（TFLite）')
    bullet(doc, '模型：康奈尔大学 BirdNET 官方 V2.4 TFLite 模型，6,522 类（全球鸟种），本地推理；')
    bullet(doc, '输入：3 秒原始波形（48kHz × 3s = 144,000 样本），float32；')
    bullet(doc, '输出：每类 logits 经 flat-sigmoid 映射为概率（官方 sensitivity=−1 / bias=1.0，即标准 sigmoid）：')
    formula(doc, r'p_k=\frac{1}{1+e^{-z_k}}')
    bullet(doc, '分窗：长录音按 3s 切窗（不足零填充），逐窗推理；')
    bullet(doc, '跨窗聚合：物种置信度取「各窗最大值」（某鸟仅在 3s 窗口鸣叫也能被捕获）；freq（出现频次）= 置信度 ≥ 阈值的窗口数（占据率）；')
    bullet(doc, '输出条目：{id, name(中文), latin, conf(0-1), freq, period(清晨/上午/黄昏/全天)}；阈值默认 0.50，可调 0.30–0.90；topK 默认 20。')
    h2(doc, '4.2 降级引擎：启发式鸟鸣检测')
    para(doc, '当模型文件缺失（离线环境）时，使用自研启发式引擎，保证功能不中断：')
    bullet(doc, '音节切分：短时能量包络 + 频带能量检测，提取候选音节（起止帧）；')
    bullet(doc, '调性检测：计算音节频谱的调性（tonalness）指标，滤除宽带噪声伪音节；')
    bullet(doc, '主频匹配：以音节主频（dominant frequency）与内置鸟种声学档案（主频区间）匹配，得到候选物种与置信度；')
    bullet(doc, '同时输出鸟声活动度 activity（0-1）。')
    h2(doc, '4.3 鸟声活动度')
    para(doc, '活动度 = 检测到鸟声的有效窗口占比，作为宜居度模型的可信度输入之一。')

    # ============ 五、生态声学指数 ============
    doc.add_page_break()
    h1(doc, '五、生态声学指数')
    para(doc, '四个指数全部基于 48kHz STFT 幅度谱计算，口径与生态声学文献一致。')
    h2(doc, '5.1 ACI —— 声学复杂度（Pieretti et al., 2011）')
    formula(doc, r'\mathrm{ACI}=\frac{1}{N_f}\sum_{f}\frac{\sum_{t}\left|dB(f,t)-dB(f,t-1)\right|}{\sum_{t}\left(dB(f,t)-\min_{f}dB(f)+1\right)}')
    para(doc, '实现：先转 dB（下限 −80dB），按频点累计相邻帧 dB 差绝对值，除以平移后 dB 和，再对频点取均值；'
              'dB 域使鸟鸣的强动态变化高于宽频噪声的小幅抖动，鸟鸣场景 ACI 更高。'
              '展示值：pct_ACI = round(100 × tanh(8.0 × ACI))。')
    h2(doc, '5.2 NDSI —— 归一化声景（Kasten et al., 2012）')
    formula(doc, r'\mathrm{NDSI}=\frac{E_{bio}-E_{anthro}}{E_{bio}+E_{anthro}},\quad \mathrm{NDSI}\in[-1,1]')
    para(doc, 'E_bio 为 2–11kHz 频带功率（生物声带），E_anthro 为 1–2kHz 频带功率（人为/交通声带）。'
              '正值表示生物声主导，负值表示人为噪声主导。展示值：pct_NDSI = (NDSI+1)/2 × 100。')
    h2(doc, '5.3 ADI —— 声学多样性（Villanueva-Rivera et al., 2011）')
    formula(doc, r'\mathrm{ADI}=-\frac{\sum_{i} p_{i}\ln p_{i}}{\ln N_{sel}}')
    para(doc, '将 0–11kHz 均分为 10 个频带；保留能量高于「最大频带能量 − 50dB」的频带（共 N_sel 个），'
              '对归一化能量 p_i 计算 Shannon 熵并除以 ln(N_sel) 归一化，∈[0,1]。')
    h2(doc, '5.4 H —— 声学熵（Sueur et al., 2008）')
    formula(doc, r'H=H_{t}\cdot H_{f},\quad H_{t}=-\frac{\sum p_{t}\ln p_{t}}{\ln N_{t}},\quad H_{f}=-\frac{\sum p_{f}\ln p_{f}}{\ln N_{f}}')
    para(doc, '时间熵 H_t：10ms 短时 RMS 包络归一化后取 Shannon 熵；频谱熵 H_f：平均幅度谱归一化后取 Shannon 熵。'
              '乘积 ∈[0,1]，越接近 1 表示各声源分布越均衡。')
    h2(doc, '5.5 频带功率')
    formula(doc, r'E(f_1,f_2)=\sum_{f\in[f_1,f_2]}\sum_{t}\left|X(t,f)\right|^{2}')
    para(doc, '作为 NDSI、ADI、噪声估计共用的基础量。')

    # ============ 六、噪声 ============
    h1(doc, '六、人为噪声占比估计')
    formula(doc, r'\mathrm{noise}=\left[\,0.65\,\frac{E_{anthro}}{E_{anthro}+E_{bio}}+0.35\,\left(1-\mathrm{activity}\right)\right]\times 100\%')
    para(doc, '双信号融合：① 频谱比（1–2kHz 人为声带 vs 2–11kHz 生物声带的能量占比）；'
              '② 鸟声活动度（有鸟声 → 噪声低）。兼顾物理声谱与语义证据。')

    # ============ 七、宜居度模型 ============
    doc.add_page_break()
    h1(doc, '七、宜居度耦合模型')
    para(doc, '宜居度（0–100）由生物多样性维度 bio 与声环境质量维度 sound 加权合成：')
    formula(doc, r'\mathrm{score}=0.55\,\mathrm{bio}+0.45\,\mathrm{sound}')
    h2(doc, '7.1 生物多样性维度 bio')
    formula(doc, r'\mathrm{bio}=18+16\,\ln(1+n_{sp})+22\,\mathrm{ADI}+7\,\overline{conf}+3\,\mathrm{activity}-0.12\,\mathrm{noise}')
    para(doc, 'n_sp：识别物种数（对数饱和）；ADI：声学多样性；conf̄：物种平均置信度；'
              'activity：鸟声活动度；noise：噪声占比（噪声掩蔽鸟声、压缩栖息空间，线性折减）。截断至 [0,100]。')
    h2(doc, '7.2 声环境质量维度 sound')
    formula(doc, r'\mathrm{sound}=10+30\,\frac{\mathrm{NDSI}+1}{2}+22\,H+18\,\left(1-\frac{\mathrm{noise}}{100}\right)')
    para(doc, 'NDSI（归一化声景）、H（声学熵）、noise（噪声）三项耦合，截断至 [0,100]。')
    h2(doc, '7.3 等级阈值')
    para(doc, 'score ≥ 70 → 宜居（Good）；50 ≤ score < 70 → 一般（Moderate）；score < 50 → 受压（Stressed）。')
    h2(doc, '7.4 置信度合成')
    formula(doc, r'\mathrm{confidence}=0.35\,\mathrm{activity}+0.30\,\overline{conf}+0.20\,\min\!\left(1,\frac{n_{sp}}{3}\right)+0.15\,\min\!\left(1,\frac{dur}{30}\right)')
    para(doc, '四路输入：鸟声活动度、识别平均置信度、物种证据充分度（3 种封顶）、采样时长充分度（30s 封顶）。'
              '置信度标签：≥0.6 高 / ≥0.4 中 / <0.4 低。')

    # ============ 八、热力图 ============
    h1(doc, '八、时间热力图')
    para(doc, '将 STFT 幅度谱按时间 × 频段分箱映射为 4×12 网格（4 个频段 × 12 个时间片），'
              '颜色编码能量强度，供用户直观理解「何时、哪个频段有鸟声」；'
              '该矩阵同样存储于分析结果，作为展示型数据（非科学计算主路径）。')

    # ============ 九、公共地图 ============
    doc.add_page_break()
    h1(doc, '九、公共地图系统')
    h2(doc, '9.1 聚合键 cluster_key')
    formula(doc, r'\mathrm{cluster\_key}=\mathrm{region\_name}\mid \mathrm{grid\_hash}(lat,\,lng)')
    para(doc, '同一地区名 + 相近坐标（网格哈希）的记录自动聚合成一个聚合点；'
              '同点多次采样保留为多条样本（支持趋势分析）。')
    h2(doc, '9.2 加权聚合')
    formula(doc, r'\mathrm{score}=\frac{\sum_{i}\mathrm{score}_{i}\cdot\max(\mathrm{conf}_{i},\,\varepsilon)}{\sum_{i}\max(\mathrm{conf}_{i},\,\varepsilon)},\quad \varepsilon=0.01')
    para(doc, '以置信度为权重的加权宜居度均值（低置信度记录权重被压缩）；'
              'confidenceAvg 为算术均值；scoreMin/scoreMax 取样本极值；n 为样本数。')
    h2(doc, '9.3 确定性坐标模糊（隐私）')
    formula(doc, r'\Delta=\mathrm{sha256}(\mathrm{key})[0:8]\Rightarrow r_{1},r_{2}\in[0,1)')
    formula(doc, r"lat'=lat+(r_{1}-0.5)\cdot 2J,\quad lng'=lng+(r_{2}-0.5)\cdot 2J,\quad J\approx \pm250m")
    para(doc, '同一 key 每次结果一致（刷新不抖动）；模糊量 ≤ ±250m，精确采样位置不可逆。')
    h2(doc, '9.4 隐私设计清单')
    bullet(doc, '公共接口不返回精确坐标（仅模糊后 lat/lng）与用户 ID；')
    bullet(doc, '样本仅公开：匿名昵称 + 日期（到天）+ 评分 + 噪声 + 鸟种；')
    bullet(doc, '删除需登录且本人校验（user_id 匹配），他人 403。')

    # ============ 十、统计与看板 ============
    h1(doc, '十、数据看板与统计')
    para(doc, 'GET /api/public/stats 聚合输出：')
    bullet(doc, 'totalSamples（总样本）、totalClusters（聚合点数）、activeUsers（活跃用户）、totalSpecies（识别鸟种数）；')
    bullet(doc, 'scoreAvg：置信度加权宜居度均值（与聚合口径一致）；')
    bullet(doc, 'buckets：宜居/一般/受压三档计数；')
    bullet(doc, 'speciesTop：跨样本物种出现次数 Top10；regionTop：按地区汇总样本数 Top10。')
    para(doc, '物种分布筛选：clusters 接口支持 species 参数，只聚合含该物种的记录（分布热力）。')

    # ============ 十一、生态简报 ============
    h1(doc, '十一、生态简报')
    para(doc, 'GET /api/public/clusters/{id}/report 生成地区生态报告：')
    bullet(doc, 'LLM 模式：配置 LLM_API_KEY（默认智谱 glm-4-flash，可换通义/DeepSeek），'
                '将聚合统计（评分/噪声/物种 Top/趋势）注入提示词，返回 Markdown 报告；')
    bullet(doc, '模板降级：无 Key 或调用失败 → 规则模板自动生成（含地区名、评分、三档结论、物种 Top 等），功能永不阻塞。')

    # ============ 十二、账号与云同步 ============
    h1(doc, '十二、账号与云同步')
    bullet(doc, '鉴权：注册/登录发放 Bearer Token（auth_tokens 表），请求头 Authorization 校验；')
    bullet(doc, '修改密码：改密后删除该用户全部 Token（强制下线所有设备）；')
    bullet(doc, '云备份：user_backups 表（user_id 主键单行），POST/GET /api/sync/backup，payload ≤2MB，账号间隔离；')
    bullet(doc, '本地数据范围：history（分析历史）、regions（地区记录）、analysis 快照、批次结果；')
    bullet(doc, 'recordedAt 回填：上传接口支持可选 recordedAt（演示/数据迁移），App 正常上传不传（服务端当前时间）。')

    # ============ 十三、前端 App ============
    doc.add_page_break()
    h1(doc, '十三、前端 App 架构')
    h2(doc, '13.1 技术结构')
    bullet(doc, 'React 18 + Vite 5 + Tailwind CSS，组件化开发，Android Kotlin WebView 套壳（含文件选择、录音权限）；')
    bullet(doc, '状态管理：Context + useReducer（appStore），支持 BACK/GO 路由、Toast、持久化；')
    bullet(doc, '数据层：repository 统一出口，mockData 与 apiService 双实现，VITE_USE_MOCK 切换演示/真实模式；')
    bullet(doc, '本地存储：localStorage（history/regions/batches），启动水合避免覆盖（skipFirstPersist）。')
    h2(doc, '13.2 数据契约守护')
    para(doc, 'frontend/tests/dataContract.test.js 守护字段契约（species/indices/livability/heatmap/mapPoints/history），'
              '前端 mock 与后端响应结构不一致时测试立即失败，保证双数据源口径一致。')
    h2(doc, '13.3 页面清单')
    para(doc, '首页 / 录音（实时+导入）/ 分析中 / 结果 / 物种清单 / 声学指数 / 时间热力图 / 宜居度详情 / '
              '地区记录 / 地图（选点）/ 鸟种图鉴（121 种）/ 登录 / 设置 / 使用说明 / 样例音频 / 方法学。')

    # ============ 十四、公共网页 ============
    h1(doc, '十四、公共地图网页')
    bullet(doc, 'MapLibre GL 渲染，高德栅格瓦片经后端 /api/tiles/* 代理（避免跨域与 Key 暴露）；')
    bullet(doc, '聚合点 circle 层：颜色按宜居度渐变（<50 红、50-70 琥珀、≥70 绿），半径按 sqrt(n) 缩放；')
    bullet(doc, '面板体系：图鉴 / 对比 / 趋势 / 简报 / 看板 / Top10 / 分享；')
    bullet(doc, '导出：聚合 CSV + 明细 CSV（UTF-8 BOM，8 列全维度）；')
    bullet(doc, '移动端：@media(max-width:700px) 全量自适应。')

    # ============ 十五、部署运维 ============
    doc.add_page_break()
    h1(doc, '十五、部署与运维')
    bullet(doc, '生产环境：阿里云轻量应用服务器（2C2G，Alibaba Cloud Linux 3），容器化部署（podman-docker 兼容层）；')
    bullet(doc, '数据持久化：Docker 命名卷 /app/data（SQLite 文件 + 模型），容器重建数据不丢；')
    bullet(doc, '更新机制：服务器 /root/update.sh（备份 db → 拉新镜像 → 重建容器 → healthz 校验）；')
    bullet(doc, 'HTTPS：nginx 反代 8000（域名 caurmosi.top，免费证书自动续期）；')
    bullet(doc, '镜像构建：GitHub Actions 推送 ghcr.io，后端镜像 + Android APK 自动打包；')
    bullet(doc, '多后端并存：Sealos（备份/过渡）与阿里云（生产）双后端，网页 API_BASE 可配置。')

    # ============ 十六、测试 ============
    h1(doc, '十六、测试与质量保障')
    three_line_table(doc,
                     ['层级', '工具', '规模', '覆盖'],
                     [['后端单元/集成', 'pytest + FastAPI TestClient', '124 用例全绿', '识别、指数、宜居度、聚合、隐私、统计、同步、鉴权'],
                      ['前端单元', 'node --test', '347 用例全绿', 'reducer、契约、联动、导出、本地化'],
                      ['数据契约', 'dataContract.test.js', '守护字段结构', 'mock 与真实 API 口径一致'],
                      ['算法校准', 'test_indices.py', '锚点回归', '好/差样地评分区间'],
                      ['E2E 验证', 'curl + 公网实测', '部署后', 'indices/tiles/login/clusters/export 全通']],
                     widths=[3.0, 3.6, 2.4, 3.4])

    # ============ 十七、参考文献 ============
    doc.add_page_break()
    h1(doc, '十七、参考文献')
    refs = [
        'Pieretti, N., Farina, A., Morri, D. (2011). A new methodology to infer the singing activity of an avian community: The Acoustic Complexity Index (ACI). Ecological Indicators, 11(3), 868-873.',
        "Kasten, E.P., Gage, S.H., Fox, J., Joo, W. (2012). The remote environmental assessment laboratory's acoustic library: An archive for studying soundscape ecology. Ecological Informatics, 12, 50-67.",
        'Villanueva-Rivera, L.J., Pijanowski, B.C., Doucette, J., Pekin, B. (2011). A primer of acoustic analysis for landscape ecologists. Landscape Ecology, 26(9), 1233-1246.',
        'Sueur, J., Pavoine, S., Hamerlynck, O., Duvail, S. (2008). Rapid acoustic survey for biodiversity appraisal. PLoS ONE, 3(12), e4065.',
        'Kahl, S., Wood, C.M., Eibl, M., Klinck, H. (2021). BirdNET: A deep learning solution for avian diversity monitoring. Ecological Informatics, 61, 101236.',
        'Klinck, H., Kahl, S. (2023). BirdNET V2.4 model & documentation (TFLite). Cornell Lab of Ornithology / Chemnitz University of Technology.',
        'Pijanowski, B.C., et al. (2011). Soundscape ecology: The science of sound in the landscape. BioScience, 61(3), 203-216.',
    ]
    for i, ref in enumerate(refs, 1):
        para(doc, f'[{i}] {ref}', indent=0)

    doc.save('F:/Desktop/linjianhuixiang-prototype/docs/《林间回响》技术设计说明书.docx')
    print('设计说明书已生成（论文风 + 数学公式）')


if __name__ == '__main__':
    build()
